import time
from docx import Document
import re
from docx.shared import Pt
from docx.oxml.ns import qn
from files import destFilePath, sourceFilePath
from file_length import FileLength


class WordHandler(object):
    regx = re.compile(r".*第\w+章.*")

    fileEncoding = ""

    # 声明Word对象
    document = Document()
    # 设置全局字体宋体，大小12
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(12)
    # 设置行间距1.5倍
    document.styles['Normal'].paragraph_format.line_spacing = 1.5

    def __init__(self):
        fileLength = FileLength()
        self.file_length = fileLength(sourceFilePath)
        self.fileEncoding = fileLength.fileEncoding
        self.start_time = time.time()

    def read_source_file(self):
        with open(sourceFilePath, 'r', encoding=self.fileEncoding) as f:
            while True:
                line = f.readline()
                if not line:
                    break
                yield line

    def write_to_word(self):
        paragraph = self.document.add_paragraph()  # 添加段落
        length = 0
        source_file = self.read_source_file()
        while True:
            try:
                line = source_file.__next__()
                length += len(line)
                if re.match(self.regx, line):
                    # 添加一级标题，设置字体为宋体
                    h_run = self.document.add_heading(level=1).add_run(line)
                    h_run.font.name = u'宋体'
                    h_run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    paragraph = self.document.add_paragraph()  # 添加段落
                elif len(line) <= 1:
                    pass
                elif paragraph is not None:
                    try:
                        paragraph.add_run(line)
                    except ValueError:
                        pass
                percentage = length * 100 / self.file_length
                print('\r' + '[完成进度]: %s%.2f%%' % ('+' * int(percentage),
                                                       float(percentage)), end='')

            except StopIteration:
                break

        self.document.save(destFilePath)
        print("\n--- 用时: %.2fs ---" % (time.time() - self.start_time))

    def test(self):
        size = 0
        source_file = self.read_source_file()
        while True:
            try:
                line = source_file.__next__()
                size += len(line)
            except StopIteration:
                break
        print(size)

    def main(self):
        # self.test()
        self.write_to_word()


if __name__ == '__main__':
    WordHandler().main()
